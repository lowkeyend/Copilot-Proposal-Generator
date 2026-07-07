from __future__ import annotations

import io

from docx import Document

from app.models.schemas import SectionResult
from app.services.docx_service import DocxComposer
from app.services.knowledge_ingest_service import _chunk_docx_document


def test_docx_ingest_uses_heading_sections() -> None:
    doc = Document()
    doc.add_heading("Scope of Work", level=1)
    doc.add_paragraph("This section covers upgrade analysis and environment preparation.")
    doc.add_heading("Testing", level=1)
    doc.add_paragraph("This section covers SIT, UAT, and rehearsal activities.")

    buffer = io.BytesIO()
    doc.save(buffer)

    chunks = _chunk_docx_document("sample.docx", buffer.getvalue())

    assert chunks
    assert any(chunk.section == "Scope of Work" for chunk in chunks)
    assert any(chunk.section == "Testing" for chunk in chunks)


def test_docx_render_sanitizer_removes_template_artifacts() -> None:
    composer = DocxComposer()
    cleaned = composer._sanitize_render_text(
        "<title>Solution</title>\n<paragraph>Body</paragraph>\n"
        "Word TOC field: Open in Word and refresh fields to populate page numbers.\n"
        "{{SECTION:Solution}}"
    )

    assert "<title>" not in cleaned.lower()
    assert "<paragraph>" not in cleaned.lower()
    assert "word toc field" not in cleaned.lower()
    assert "{{section:solution}}" not in cleaned.lower()


def test_reference_heading_replacement_uses_matching_heading() -> None:
    doc = Document()
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph("Old summary")
    doc.add_heading("Company Profile", level=1)
    doc.add_paragraph("Static profile")

    composer = DocxComposer()
    heading = composer._find_heading(doc, "Executive Summary")
    assert heading is not None

    composer._replace_heading_section(
        doc,
        heading,
        "New grounded summary.\n\n- Point A\n- Point B",
    )

    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    assert "New grounded summary." in text
    assert "Old summary" not in text
    assert "Static profile" in text


def test_reference_heading_replacement_preserves_visual_paragraphs() -> None:
    doc = Document()
    heading = doc.add_heading("Scope of Work", level=1)
    doc.add_paragraph("Old scope text")
    visual = doc.add_paragraph("")
    visual._element.append(visual._element.makeelement("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"))
    doc.add_heading("Project Timeline", level=1)

    composer = DocxComposer()
    composer._replace_heading_section(doc, heading, "New scope text")

    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    assert "New scope text" in text
    assert "Old scope text" not in text
