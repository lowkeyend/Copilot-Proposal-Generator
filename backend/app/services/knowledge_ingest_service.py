from __future__ import annotations

import io
import re
from pathlib import Path
from dataclasses import dataclass
from typing import Iterable
from uuid import uuid4
from zipfile import ZipFile

from docx import Document
from fastapi import UploadFile

from app.config import get_settings
from app.services.qdrant_service import get_qdrant

try:
    from pypdf import PdfReader
except Exception:  # pragma: no cover
    PdfReader = None


def _clean(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def _clean_multiline(text: str) -> str:
    text = (text or "").replace("\x00", " ").replace("\r\n", "\n").replace("\r", "\n")
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.split("\n")]
    collapsed: list[str] = []
    blank_run = 0
    for line in lines:
        if not line:
            blank_run += 1
            if blank_run <= 1:
                collapsed.append("")
            continue
        blank_run = 0
        collapsed.append(line)
    return "\n".join(collapsed).strip()


def _summary(text: str, limit: int = 12) -> str:
    words = text.split()
    if not words:
        return "Untitled chunk"
    head = " ".join(words[:limit])
    return head if len(words) <= limit else f"{head}..."


def _split_sentences(text: str) -> list[str]:
    parts = re.split(r"(?<=[.!?])\s+", text)
    return [part.strip() for part in parts if part.strip()]


def _chunk_text(text: str, chunk_size: int = 420, overlap: int = 70) -> list[str]:
    sentences = _split_sentences(text)
    if not sentences:
        return [text] if text else []

    chunks: list[str] = []
    current: list[str] = []
    current_len = 0

    for sentence in sentences:
        sent_len = len(sentence)
        if current and current_len + sent_len + 1 > chunk_size:
            chunk = " ".join(current).strip()
            if chunk:
                chunks.append(chunk)
            tail: list[str] = []
            tail_len = 0
            for prior in reversed(current):
                if tail_len + len(prior) > overlap:
                    break
                tail.insert(0, prior)
                tail_len += len(prior)
            current = [*tail, sentence]
            current_len = sum(len(part) for part in current)
        else:
            current.append(sentence)
            current_len += sent_len + 1

    final = " ".join(current).strip()
    if final:
        chunks.append(final)
    return [_clean(chunk) for chunk in chunks if _clean(chunk)]


def _chunk_docx_paragraphs(text: str, chunk_size: int = 700) -> list[str]:
    paragraphs = [line.strip() for line in (text or "").splitlines() if line.strip()]
    chunks: list[str] = []
    current: list[str] = []
    size = 0
    for paragraph in paragraphs:
        if current and size + len(paragraph) + 2 > chunk_size:
            chunks.append("\n\n".join(current).strip())
            current = []
            size = 0
        current.append(paragraph)
        size += len(paragraph) + 2
    if current:
        chunks.append("\n\n".join(current).strip())
    return [chunk for chunk in chunks if chunk]


def _batched(items: list, size: int = 16):
    for index in range(0, len(items), size):
        yield items[index : index + size]


async def _read_upload(file: UploadFile) -> bytes:
    data = await file.read()
    await file.seek(0)
    return data


def _extract_docx(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    blocks = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))
    return _clean("\n".join(blocks))


def _is_heading_paragraph(paragraph) -> bool:
    style_name = ((getattr(paragraph.style, "name", "") or "").strip()).lower()
    if not style_name:
        return False
    return style_name.startswith("heading") or style_name in {"title", "subtitle"}


def _extract_docx_sections(data: bytes) -> list[tuple[str, str]]:
    doc = Document(io.BytesIO(data))
    sections: list[tuple[str, str]] = []
    current_heading = "Document Overview"
    current_blocks: list[str] = []

    def flush() -> None:
        nonlocal current_blocks
        body = _clean_multiline("\n".join(current_blocks))
        if body:
            sections.append((current_heading, body))
        current_blocks = []

    for paragraph in doc.paragraphs:
        text = _clean(paragraph.text or "")
        if not text:
            continue
        if _is_heading_paragraph(paragraph):
            flush()
            current_heading = text
            continue
        current_blocks.append(text)

    for table in doc.tables:
        current_blocks.extend(
            " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            for row in table.rows
            if any(cell.text.strip() for cell in row.cells)
        )

    flush()
    return sections or [("Document Overview", _extract_docx(data))]


def _safe_name(name: str) -> str:
    return re.sub(r"[^A-Za-z0-9._-]+", "_", name).strip("_") or "document"


def _extract_docx_images(filename: str, data: bytes) -> list[str]:
    settings = get_settings()
    base = _safe_name(Path(filename).stem)
    image_dir = settings.assets_path / "document-images" / base
    image_dir.mkdir(parents=True, exist_ok=True)

    saved: list[str] = []
    try:
        with ZipFile(io.BytesIO(data)) as archive:
            media = [
                name
                for name in archive.namelist()
                if name.startswith("word/media/")
                and not name.endswith("/")
            ]
            for idx, member in enumerate(media, start=1):
                ext = Path(member).suffix.lower() or ".bin"
                target = image_dir / f"image_{idx:02d}{ext}"
                target.write_bytes(archive.read(member))
                saved.append(str(target))
    except Exception:
        return []
    return saved


def _extract_pdf(data: bytes) -> str:
    if PdfReader is None:
        raise RuntimeError("PDF support is unavailable because pypdf is not installed.")
    reader = PdfReader(io.BytesIO(data))
    # Preserve page and line boundaries so headings, tables, and workstreams
    # remain separable during section-aware chunking.
    return _clean_multiline("\n\n".join(page.extract_text() or "" for page in reader.pages))


def _extract_text(name: str, data: bytes) -> str:
    lower = name.lower()
    if lower.endswith(".docx"):
        return _extract_docx(data)
    if lower.endswith(".pdf"):
        return _extract_pdf(data)
    if lower.endswith((".txt", ".md")):
        return _clean_multiline(data.decode("utf-8", errors="ignore"))
    raise RuntimeError("Unsupported file type. Upload .docx, .pdf, .txt, or .md files.")


def _looks_like_plaintext_heading(line: str) -> bool:
    stripped = (line or "").strip().strip("#").strip()
    if not stripped:
        return False
    if len(stripped) > 80:
        return False
    if stripped.endswith(":") and len(stripped.split()) <= 8:
        return True
    alpha = re.sub(r"[^A-Za-z ]+", "", stripped)
    words = [word for word in alpha.split() if word]
    if 0 < len(words) <= 8 and sum(1 for word in words if word[:1].isupper()) >= max(1, len(words) - 1):
        return True
    if stripped.isupper() and 0 < len(stripped.split()) <= 10:
        return True
    return False


def _is_tabular_line(line: str) -> bool:
    stripped = (line or "").strip()
    if not stripped:
        return False
    if "|" in stripped:
        return True
    if re.search(r"\bmandays?\b|\brate\b|\bcommercial quote\b|\$\s*\d", stripped, flags=re.IGNORECASE):
        return True
    if len(re.findall(r"\d+", stripped)) >= 4:
        return True
    return False


def _extract_plaintext_sections(name: str, text: str) -> list[tuple[str, list[str]]]:
    lines = [line.rstrip() for line in (text or "").split("\n")]
    sections: list[tuple[str, list[str]]] = []
    current_heading = Path(name).stem or "Document Overview"
    current_lines: list[str] = []

    def flush() -> None:
        nonlocal current_lines
        cleaned = [line for line in current_lines if line.strip()]
        if cleaned:
            sections.append((current_heading, cleaned))
        current_lines = []

    for raw in lines:
        line = raw.strip()
        if not line:
            current_lines.append("")
            continue
        if _looks_like_plaintext_heading(line):
            flush()
            current_heading = line.strip("# ").strip()
            continue
        current_lines.append(line)

    flush()
    return sections or [(Path(name).stem or "Document Overview", [line for line in lines if line.strip()])]


def _chunk_generic_document(name: str, text: str) -> list[ParsedChunk]:
    parsed: list[ParsedChunk] = []
    for heading, lines in _extract_plaintext_sections(name, text):
        block: list[str] = []
        block_is_table = False

        def flush_block() -> None:
            nonlocal block, block_is_table
            joined = "\n".join(block if block_is_table else [item for item in block if item.strip()]).strip()
            if not joined:
                block = []
                block_is_table = False
                return
            chunks = [joined] if block_is_table else _chunk_text(joined)
            for chunk in chunks:
                parsed.append(
                    ParsedChunk(
                        text=chunk,
                        section=heading or "Document Overview",
                        summary=f"{heading or 'Document Overview'}: {_summary(_clean(chunk), limit=10)}",
                    )
                )
            block = []
            block_is_table = False

        for line in lines:
            if not line.strip():
                flush_block()
                continue
            line_is_table = _is_tabular_line(line)
            if block and line_is_table != block_is_table:
                flush_block()
            block.append(line)
            block_is_table = line_is_table

        flush_block()

    return parsed or [
        ParsedChunk(
            text=chunk,
            section=Path(name).stem or "Document Overview",
            summary=f"{Path(name).stem or 'Document Overview'}: {_summary(chunk, limit=10)}",
        )
        for chunk in _chunk_text(_clean(text))
    ]


def _chunk_docx_document(name: str, data: bytes) -> list[ParsedChunk]:
    parsed: list[ParsedChunk] = []
    for heading, body in _extract_docx_sections(data):
        for chunk in _chunk_docx_paragraphs(body):
            parsed.append(
                ParsedChunk(
                    text=chunk,
                    section=heading or "Document Overview",
                    summary=f"{heading or 'Document Overview'}: {_summary(chunk, limit=10)}",
                )
            )
    return parsed or _chunk_generic_document(name, _extract_docx(data))


@dataclass
class ParsedDocument:
    filename: str
    text: str
    image_paths: list[str]
    raw_bytes: bytes


@dataclass
class ParsedChunk:
    text: str
    section: str
    summary: str


class KnowledgeIngestService:
    async def parse_files(self, files: Iterable[UploadFile]) -> list[ParsedDocument]:
        parsed: list[ParsedDocument] = []
        for file in files:
            name = file.filename or "document"
            data = await _read_upload(file)
            text = _extract_text(name, data)
            if not text:
                continue
            images = _extract_docx_images(name, data) if name.lower().endswith(".docx") else []
            parsed.append(ParsedDocument(filename=name, text=text, image_paths=images, raw_bytes=data))
        return parsed

    async def ingest_files(
        self,
        files: list[UploadFile],
        source_proposal: str,
        source_section: str,
        proposal_family: str,
    ) -> tuple[list[str], int]:
        qdrant = get_qdrant()
        filenames: list[str] = []
        total_points = 0

        # Flush each upload in small batches so large DOCX files do not exhaust
        # the memory available on a hosted backend.
        for file in files:
            name = file.filename or "document"
            data = await _read_upload(file)
            text = _extract_text(name, data)
            if not text:
                continue
            filenames.append(name)
            image_paths = _extract_docx_images(name, data) if name.lower().endswith(".docx") else []
            chunks = (
                _chunk_docx_document(name, data)
                if name.lower().endswith(".docx")
                else _chunk_generic_document(name, text)
            )
            batch = []
            for index, chunk in enumerate(chunks, start=1):
                payload = {
                    "text": chunk.text,
                    "chunk_text": chunk.text,
                    "chunk_summary": chunk.summary,
                    "source_proposal": source_proposal or name,
                    "source_section": source_section or chunk.section or f"Upload chunk {index}",
                    "proposal_family": proposal_family or "Uploaded Knowledge",
                    "file": name,
                    "document_name": name,
                    "section": source_section or chunk.section or f"Upload chunk {index}",
                    "chunk_index": index,
                    "image_paths": image_paths,
                }
                batch.append(
                    qdrant.build_point(
                        chunk_id=uuid4().hex,
                        text=chunk.text,
                        payload=payload,
                    )
                )
                if len(batch) >= 8:
                    qdrant.upsert_points(batch)
                    total_points += len(batch)
                    batch = []
            if batch:
                qdrant.upsert_points(batch)
                total_points += len(batch)
        if not filenames:
            raise RuntimeError("No readable content found in the uploaded files.")
        return filenames, total_points


_ingest_singleton: KnowledgeIngestService | None = None


def get_knowledge_ingest() -> KnowledgeIngestService:
    global _ingest_singleton
    if _ingest_singleton is None:
        _ingest_singleton = KnowledgeIngestService()
    return _ingest_singleton
