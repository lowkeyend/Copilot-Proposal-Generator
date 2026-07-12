"""Agent 9 — DOCX Composer.

Renders a reviewed proposal into a professional, branded Word document:
title page, auto-updating Table of Contents field, heading hierarchy,
markdown-ish body parsing (headings / bullets / simple pipe tables),
running header, and footer with page numbering.

Uses python-docx plus a few low-level OOXML fields (TOC + PAGE) that
python-docx doesn't expose directly.
"""

from __future__ import annotations

import re
import urllib.request
from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph

from app.config import get_settings
from app.models.schemas import ClientContext, SectionResult, TemplateBlock

BRAND = RGBColor(0x1F, 0x3A, 0x5F)  # deep navy
ACCENT = RGBColor(0x2E, 0x6F, 0x8E)
BODY_FONT = "Century Gothic"
STATIC_TEMPLATE_HEADINGS = {"Company Profile", "Case Studies"}
SECTION_HEADING_ALIASES = {
    "Solution": "Proposed Solution",
    "Methodology": "Upgrade Methodology",
    "Governance": "Project Governance",
    "Training": "Upgrade Methodology",
}


def _set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold


def _add_field(paragraph, instruction: str) -> None:
    """Insert a Word field (e.g. TOC, PAGE) into a paragraph."""
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def _clear_paragraph(paragraph) -> None:
    element = paragraph._element
    for child in list(element):
        element.remove(child)


def _paragraph_after(paragraph, style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    created = Paragraph(new_p, paragraph._parent)
    if style:
        try:
            created.style = style
        except KeyError:
            pass
    return created


def _element_has_visual(element) -> bool:
    xml = element.xml if hasattr(element, "xml") else ""
    return any(marker in xml for marker in ("w:drawing", "w:pict", "pic:pic", "v:shape"))


class DocxComposer:
    def __init__(self) -> None:
        self.settings = get_settings()

    # ------------------------------------------------------------------
    def compose(
        self,
        title: str,
        context: ClientContext,
        sections: list[SectionResult],
        proposal_id: Optional[str] = None,
    ) -> Path:
        template_path = self._ensure_template()
        if template_path.exists():
            doc = Document(str(template_path))
            self._fill_template_metadata(doc, title, context)
            self._inject_template_content(doc, sections)
        else:
            doc = Document()
            self._configure_styles(doc)
            self._add_header_footer(doc, title)
            self._add_title_page(doc, title, context)
            self._add_toc(doc)
            self._add_sections(doc, sections)

        self.settings.generated_path.mkdir(parents=True, exist_ok=True)
        safe = re.sub(r"[^A-Za-z0-9_-]+", "_", title).strip("_") or "proposal"
        stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{safe}_{stamp}.docx"
        out_path = self.settings.generated_path / filename
        doc.save(str(out_path))
        return out_path

    def _ensure_template(self) -> Path:
        template_path = self.settings.proposal_template_path
        if template_path.exists():
            return template_path
        template_url = self.settings.proposal_template_url.strip()
        if not template_url:
            return template_path
        template_path.parent.mkdir(parents=True, exist_ok=True)
        try:
            with urllib.request.urlopen(template_url, timeout=20) as response:
                template_path.write_bytes(response.read())
        except Exception:
            return template_path
        return template_path

    # ------------------------------------------------------------------
    def _configure_styles(self, doc: Document) -> None:
        for section in doc.sections:
            section.top_margin = Inches(0.6)
            section.bottom_margin = Inches(0.6)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)
        normal = doc.styles["Normal"]
        normal.font.name = BODY_FONT
        normal.font.size = Pt(10.5)
        for level, size in ((1, 16), (2, 13), (3, 11.5)):
            try:
                h = doc.styles[f"Heading {level}"]
                h.font.name = BODY_FONT
                h.font.color.rgb = BRAND
                h.font.size = Pt(size)
                h.font.bold = True
            except KeyError:
                continue
        try:
            title = doc.styles["Title"]
            title.font.name = BODY_FONT
            title.font.color.rgb = BRAND
            title.font.size = Pt(28)
            title.font.bold = True
        except KeyError:
            pass

    def _add_header_footer(self, doc: Document, title: str) -> None:
        section = doc.sections[0]

        header_p = section.header.paragraphs[0]
        header_p.text = title
        header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if header_p.runs:
            header_p.runs[0].font.size = Pt(8)
            header_p.runs[0].font.color.rgb = ACCENT

        footer_p = section.footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_p.add_run("Page ").font.size = Pt(8)
        _add_field(footer_p, "PAGE")
        footer_p.add_run(" of ").font.size = Pt(8)
        _add_field(footer_p, "NUMPAGES")

    def _add_title_page(
        self, doc: Document, title: str, context: ClientContext
    ) -> None:
        lead = doc.add_paragraph()
        lead.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lead_run = lead.add_run("PROPOSAL")
        lead_run.bold = True
        lead_run.font.size = Pt(11)
        lead_run.font.color.rgb = ACCENT
        lead_run.font.name = BODY_FONT

        for _ in range(2):
            doc.add_paragraph()
        t = doc.add_paragraph()
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = t.add_run(title)
        run.bold = True
        run.font.size = Pt(28)
        run.font.color.rgb = BRAND
        run.font.name = BODY_FONT

        if context.client_name:
            sub = doc.add_paragraph()
            sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = sub.add_run(f"Prepared for {context.client_name}")
            r.font.size = Pt(15)
            r.font.color.rgb = ACCENT
            r.font.name = BODY_FONT

        strap = doc.add_paragraph()
        strap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        strap_run = strap.add_run(
            "Strategic delivery plan, implementation approach, and commercial alignment"
        )
        strap_run.font.size = Pt(10.5)
        strap_run.italic = True
        strap_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        strap_run.font.name = BODY_FONT

        doc.add_paragraph()
        summary = doc.add_table(rows=0, cols=2)
        summary.style = "Light Grid Accent 1"
        summary.autofit = True
        meta_lines = [
            ("Industry", context.industry),
            ("Engagement", context.project_type),
            ("Date", datetime.now().strftime("%d %B %Y")),
        ]
        for label, value in meta_lines:
            if not value:
                continue
            row = summary.add_row().cells
            _set_cell_text(row[0], label, bold=True)
            _set_cell_text(row[1], value)
        doc.add_paragraph()
        note = doc.add_paragraph()
        note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        note_run = note.add_run("Confidential draft prepared for review")
        note_run.font.size = Pt(9)
        note_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        note_run.font.name = BODY_FONT
        doc.add_page_break()

    def _add_toc(self, doc: Document) -> None:
        heading = doc.add_paragraph("Table of Contents")
        heading.style = doc.styles["Heading 1"]
        field_p = doc.add_paragraph()
        _add_field(field_p, 'TOC \\o "1-3" \\h \\z \\u')
        doc.add_page_break()

    def _add_sections(self, doc: Document, sections: list[SectionResult]) -> None:
        for idx, section in enumerate(sections, 1):
            doc.add_heading(f"{idx}. {section.title}", level=1)
            if getattr(section, "blocks", None):
                self._render_blocks(doc, section.blocks)
            else:
                self._render_markdownish(doc, section.content)
                self._render_evidence_images(doc, section)
            doc.add_paragraph()

    def _render_evidence_images(self, doc: Document, section: SectionResult) -> None:
        image_paths: list[str] = []
        for image in getattr(section, "images", []) or []:
            raw = getattr(image, "asset_path", "") or getattr(image, "asset_url", "")
            if raw and raw not in image_paths:
                image_paths.append(raw)
        for evidence in section.evidence or []:
            for path in getattr(evidence, "image_paths", []) or []:
                if path not in image_paths:
                    image_paths.append(path)
        if not image_paths:
            return

        added = 0
        for raw_path in image_paths:
            if added >= 2:
                break
            path = Path(raw_path)
            if not path.is_absolute():
                path = self.settings.assets_path / path
            if not path.exists():
                continue
            if path.suffix.lower() not in {".png", ".jpg", ".jpeg", ".bmp", ".gif", ".tiff"}:
                continue
            try:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(str(path), width=Inches(5.9))
                caption = doc.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = caption.add_run(f"Figure {added + 1}: {section.title}")
                run.italic = True
                run.font.size = Pt(8.5)
                run.font.name = BODY_FONT
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                added += 1
            except Exception:
                continue

    def _find_paragraph(self, doc: Document, token: str):
        target = token.strip().lower()
        for paragraph in doc.paragraphs:
            if target in (paragraph.text or "").strip().lower():
                return paragraph
        return None

    def _find_heading(self, doc: Document, title: str):
        wanted = {
            title.strip().lower(),
            SECTION_HEADING_ALIASES.get(title.strip(), "").strip().lower(),
        }
        wanted = {item for item in wanted if item}
        for paragraph in doc.paragraphs:
            style_name = (getattr(getattr(paragraph, "style", None), "name", "") or "").lower()
            text = (paragraph.text or "").strip().lower()
            if text in wanted and ("heading" in style_name or style_name.startswith("cn head")):
                return paragraph
        return None

    def _paragraph_heading_level(self, paragraph) -> int:
        style_name = (getattr(getattr(paragraph, "style", None), "name", "") or "").lower()
        match = re.search(r"heading\s+(\d+)", style_name)
        if match:
            return int(match.group(1))
        if "cn head 1" in style_name:
            return 1
        return 9

    def _delete_paragraph(self, paragraph) -> None:
        element = paragraph._element
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)

    def _replace_heading_section(self, doc: Document, heading, content: str) -> None:
        current_level = self._paragraph_heading_level(heading)
        cursor = heading._element.getnext()
        removable = []

        while cursor is not None:
            next_cursor = cursor.getnext()
            paragraph = next(
                (item for item in doc.paragraphs if item._element == cursor),
                None,
            )
            if paragraph is not None:
                style_name = (getattr(getattr(paragraph, "style", None), "name", "") or "").lower()
                if "heading" in style_name or style_name.startswith("cn head"):
                    next_level = self._paragraph_heading_level(paragraph)
                    if next_level <= current_level:
                        break
                if not _element_has_visual(cursor):
                    removable.append(cursor)
            elif cursor.tag.endswith("}tbl"):
                if not _element_has_visual(cursor):
                    removable.append(cursor)
            else:
                if not _element_has_visual(cursor):
                    removable.append(cursor)
            cursor = next_cursor

        for element in removable:
            parent = element.getparent()
            if parent is not None:
                parent.remove(element)

        self._render_markdownish_after(heading, content)

    def _fill_template_metadata(self, doc: Document, title: str, context: ClientContext) -> None:
        replacements = {
            "{{PROPOSAL_TITLE}}": title,
            "{{CLIENT_NAME}}": context.client_name or "",
            "{{INDUSTRY}}": context.industry or "",
            "{{PROJECT_TYPE}}": context.project_type or "",
            "{{DATE}}": datetime.now().strftime("%d %B %Y"),
        }
        for paragraph in doc.paragraphs:
            text = paragraph.text or ""
            new_text = text
            for token, value in replacements.items():
                new_text = new_text.replace(token, value)
            if new_text != text:
                paragraph.text = new_text

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text or ""
                        new_text = text
                        for token, value in replacements.items():
                            new_text = new_text.replace(token, value)
                        if new_text != text:
                            paragraph.text = new_text

    def _inject_toc_at_placeholder(self, doc: Document) -> bool:
        paragraph = self._find_paragraph(doc, "{{TOC}}")
        if paragraph is None:
            return False
        _clear_paragraph(paragraph)
        _add_field(paragraph, 'TOC \\o "1-3" \\h \\z \\u')
        return True

    def _inject_template_content(self, doc: Document, sections: list[SectionResult]) -> None:
        self._inject_toc_at_placeholder(doc)
        inserted_any = False
        for section in sections:
            if section.title in STATIC_TEMPLATE_HEADINGS:
                continue
            token = f"{{{{SECTION:{section.title}}}}}".lower()
            paragraph = self._find_paragraph(doc, token)
            if paragraph is None:
                heading = self._find_heading(doc, section.title)
                if heading is not None:
                    if getattr(section, "blocks", None):
                        self._replace_heading_section_with_blocks(doc, heading, section.blocks)
                    else:
                        self._replace_heading_section(doc, heading, section.content)
                        self._render_evidence_images(doc, section)
                    inserted_any = True
                continue
            _clear_paragraph(paragraph)
            if getattr(section, "blocks", None):
                self._render_blocks_into_paragraph(doc, paragraph, section.blocks)
            else:
                self._render_markdownish_into_paragraph(doc, paragraph, section.content)
                self._render_evidence_images(doc, section)
            inserted_any = True

        collection_token = self._find_paragraph(doc, "{{SECTIONS}}")
        if collection_token is not None:
            _clear_paragraph(collection_token)
            for section in sections:
                doc.add_heading(section.title, level=1)
                if getattr(section, "blocks", None):
                    self._render_blocks(doc, section.blocks)
                else:
                    self._render_markdownish(doc, section.content)
                    self._render_evidence_images(doc, section)
                doc.add_paragraph()
            inserted_any = True

        if not inserted_any:
            if not self._inject_toc_at_placeholder(doc):
                self._add_toc(doc)
            self._add_sections(doc, sections)

    def _render_markdownish_into_paragraph(self, doc: Document, paragraph, content: str) -> None:
        sanitized = self._sanitize_render_text(content)
        blocks = [block.strip() for block in re.split(r"\n{2,}", sanitized) if block.strip()]
        if not blocks:
            return
        first, *rest = blocks
        self._add_inline(paragraph, first)
        for block in rest:
            p = doc.add_paragraph()
            self._add_inline(p, block)

    def _render_blocks_into_paragraph(self, doc: Document, paragraph, blocks: list[TemplateBlock]) -> None:
        if not blocks:
            return
        first_text = next((block.text for block in blocks if block.kind in {"heading", "paragraph"} and block.text.strip()), "")
        if first_text:
            self._add_inline(paragraph, first_text)
        cursor = paragraph
        for block in blocks[1:]:
            cursor = self._render_block_after(cursor, block)

    def _replace_heading_section_with_blocks(self, doc: Document, heading, blocks: list[TemplateBlock]) -> None:
        current_level = self._paragraph_heading_level(heading)
        cursor = heading._element.getnext()
        removable = []

        while cursor is not None:
            next_cursor = cursor.getnext()
            paragraph = next((item for item in doc.paragraphs if item._element == cursor), None)
            if paragraph is not None:
                style_name = (getattr(getattr(paragraph, "style", None), "name", "") or "").lower()
                if "heading" in style_name or style_name.startswith("cn head"):
                    next_level = self._paragraph_heading_level(paragraph)
                    if next_level <= current_level:
                        break
                if not _element_has_visual(cursor):
                    removable.append(cursor)
            elif cursor.tag.endswith("}tbl"):
                if not _element_has_visual(cursor):
                    removable.append(cursor)
            else:
                if not _element_has_visual(cursor):
                    removable.append(cursor)
            cursor = next_cursor
        for element in removable:
            parent = element.getparent()
            if parent is not None:
                parent.remove(element)
        cursor_p = heading
        for block in blocks:
            if block.kind == "heading":
                if block.heading_level <= 1:
                    heading.text = block.text
                    cursor_p = heading
                    continue
            cursor_p = self._render_block_after(cursor_p, block)

    def _render_block_after(self, cursor: Paragraph, block: TemplateBlock) -> Paragraph:
        if block.kind == "heading":
            style = "Heading 2" if block.heading_level <= 2 else "Heading 3"
            created = _paragraph_after(cursor, style=style)
            self._add_inline(created, block.text)
            return created
        if block.kind == "paragraph":
            created = _paragraph_after(cursor)
            self._add_inline(created, block.text)
            return created
        if block.kind == "list":
            created = cursor
            items = block.items or [line.strip() for line in block.text.splitlines() if line.strip()]
            for item in items:
                created = _paragraph_after(created, style="List Bullet")
                self._add_inline(created, item)
            return created
        if block.kind == "table":
            created = _paragraph_after(cursor)
            rows = block.table_rows or []
            if not rows:
                return created
            cols = max(len(row) for row in rows)
            table = created._parent.add_table(rows=0, cols=cols)
            table._element.getparent().remove(table._element)
            created._p.addnext(table._element)
            table.style = "Light Grid Accent 1"
            for ridx, row in enumerate(rows):
                cells = table.add_row().cells
                for cidx in range(cols):
                    val = row[cidx] if cidx < len(row) else ""
                    _set_cell_text(cells[cidx], val, bold=(ridx == 0))
            return created
        if block.kind == "image" and block.image:
            created = _paragraph_after(cursor)
            path = Path(block.image.asset_path or "")
            if not path.is_absolute():
                path = self.settings.assets_path / path
            if path.exists():
                created.alignment = WD_ALIGN_PARAGRAPH.CENTER
                created.add_run().add_picture(str(path), width=Inches(5.9))
                caption = _paragraph_after(created)
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = caption.add_run(block.image.caption or block.image.filename or "")
                run.italic = True
                run.font.size = Pt(8.5)
                run.font.name = BODY_FONT
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                return caption
            return created
        return cursor

    def _render_blocks(self, doc: Document, blocks: list[TemplateBlock]) -> None:
        for block in blocks:
            if block.kind == "heading":
                level = 2 if block.heading_level <= 2 else 3
                doc.add_heading(block.text, level=level)
            elif block.kind == "paragraph":
                p = doc.add_paragraph()
                self._add_inline(p, block.text)
            elif block.kind == "list":
                for item in block.items or [line.strip() for line in block.text.splitlines() if line.strip()]:
                    p = doc.add_paragraph(style="List Bullet")
                    self._add_inline(p, item)
            elif block.kind == "table":
                rows = block.table_rows or []
                if not rows:
                    continue
                table = doc.add_table(rows=0, cols=max(len(row) for row in rows))
                table.style = "Light Grid Accent 1"
                for ridx, row in enumerate(rows):
                    cells = table.add_row().cells
                    for cidx in range(len(cells)):
                        val = row[cidx] if cidx < len(row) else ""
                        _set_cell_text(cells[cidx], val, bold=(ridx == 0))
            elif block.kind == "image" and block.image:
                self._render_block_after(doc.add_paragraph(), block)

    def _render_markdownish_after(self, anchor: Paragraph, content: str) -> Paragraph:
        lines = self._sanitize_render_text(content).splitlines()
        table_buffer: list[str] = []
        cursor = anchor

        def flush_table() -> None:
            nonlocal table_buffer, cursor
            rows = [r for r in table_buffer if r.strip()]
            rows = [r for r in rows if not re.match(r"^\s*\|?[\s:|-]+\|?\s*$", r)]
            if not rows:
                table_buffer = []
                return
            table_p = _paragraph_after(cursor)
            parsed = [
                [c.strip() for c in r.strip().strip("|").split("|")] for r in rows
            ]
            cols = max(len(r) for r in parsed)
            table = table_p._parent.add_table(rows=0, cols=cols)
            table._element.getparent().remove(table._element)
            table_p._p.addnext(table._element)
            for ridx, row in enumerate(parsed):
                cells = table.add_row().cells
                for cidx in range(cols):
                    val = row[cidx] if cidx < len(row) else ""
                    _set_cell_text(cells[cidx], val, bold=(ridx == 0))
            cursor = table_p
            table_buffer = []

        for raw_line in lines:
            line = raw_line.rstrip()
            if "|" in line and line.strip().startswith("|"):
                table_buffer.append(line)
                continue
            if table_buffer:
                flush_table()
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith("### "):
                cursor = _paragraph_after(cursor, style="Heading 3")
                self._add_inline(cursor, stripped[4:])
            elif stripped.startswith("## "):
                cursor = _paragraph_after(cursor, style="Heading 2")
                self._add_inline(cursor, stripped[3:])
            elif stripped.startswith("# "):
                cursor = _paragraph_after(cursor, style="Heading 2")
                self._add_inline(cursor, stripped[2:])
            elif re.match(r"^[-*]\s+", stripped):
                cursor = _paragraph_after(cursor, style="List Bullet")
                self._add_inline(cursor, re.sub(r"^[-*]\s+", "", stripped))
            elif re.match(r"^\d+[.)]\s+", stripped):
                cursor = _paragraph_after(cursor, style="List Number")
                self._add_inline(cursor, re.sub(r"^\d+[.)]\s+", "", stripped))
            else:
                cursor = _paragraph_after(cursor)
                self._add_inline(cursor, stripped)

        if table_buffer:
            flush_table()
        return cursor

    # ------------------------------------------------------------------
    def _render_markdownish(self, doc: Document, content: str) -> None:
        """Render a lightweight subset of markdown the LLM tends to produce."""
        lines = self._sanitize_render_text(content).splitlines()
        i = 0
        table_buffer: list[str] = []

        def flush_table() -> None:
            nonlocal table_buffer
            rows = [r for r in table_buffer if r.strip()]
            # Drop separator rows like |---|---|
            rows = [r for r in rows if not re.match(r"^\s*\|?[\s:|-]+\|?\s*$", r)]
            if not rows:
                table_buffer = []
                return
            parsed = [
                [c.strip() for c in r.strip().strip("|").split("|")] for r in rows
            ]
            cols = max(len(r) for r in parsed)
            table = doc.add_table(rows=0, cols=cols)
            table.style = "Light Grid Accent 1"
            for ridx, row in enumerate(parsed):
                cells = table.add_row().cells
                for cidx in range(cols):
                    val = row[cidx] if cidx < len(row) else ""
                    _set_cell_text(cells[cidx], val, bold=(ridx == 0))
            table_buffer = []

        while i < len(lines):
            line = lines[i].rstrip()
            if "|" in line and line.strip().startswith("|"):
                table_buffer.append(line)
                i += 1
                continue
            if table_buffer:
                flush_table()

            stripped = line.strip()
            if not stripped:
                i += 1
                continue
            if stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=2)
            elif re.match(r"^[-*]\s+", stripped):
                p = doc.add_paragraph(style="List Bullet")
                self._add_inline(p, re.sub(r"^[-*]\s+", "", stripped))
            elif re.match(r"^\d+[.)]\s+", stripped):
                p = doc.add_paragraph(style="List Number")
                self._add_inline(p, re.sub(r"^\d+[.)]\s+", "", stripped))
            else:
                p = doc.add_paragraph()
                self._add_inline(p, stripped)
            i += 1

        if table_buffer:
            flush_table()

    def _sanitize_render_text(self, content: str) -> str:
        cleaned = content or ""
        cleaned = re.sub(r"</?(?:title|paragraph|h1|h2|h3|body|html|xml|div|span|p)\b[^>]*>", "", cleaned, flags=re.IGNORECASE)
        cleaned = cleaned.replace("<section>", "").replace("</section>", "")
        cleaned = re.sub(r"\{\{[^}]+\}\}", "", cleaned)
        cleaned = re.sub(r"(?im)^word toc field:.*$", "", cleaned)
        cleaned = re.sub(r"(?im)^open in word and refresh fields.*$", "", cleaned)
        cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
        return cleaned.strip()

    def _add_inline(self, paragraph, text: str) -> None:
        """Handle **bold** spans inline."""
        parts = re.split(r"(\*\*[^*]+\*\*)", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                paragraph.add_run(part[2:-2]).bold = True
            elif part:
                paragraph.add_run(part)


_composer_singleton: Optional[DocxComposer] = None


def get_composer() -> DocxComposer:
    global _composer_singleton
    if _composer_singleton is None:
        _composer_singleton = DocxComposer()
    return _composer_singleton
