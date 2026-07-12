from __future__ import annotations

import re
from pathlib import Path
from typing import Any, Iterator
from uuid import uuid4

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from app.config import get_settings
from app.models.schemas import (
    TemplateBlock,
    TemplateDocumentArtifact,
    TemplateImage,
    TemplateParagraph,
    TemplateSectionNode,
    TemplateTable,
)

_STATIC_SECTIONS = {
    "company profile",
    "about systems limited",
    "about systems ltd",
    "case studies",
}


def _is_heading(style_name: str) -> bool:
    return (style_name or "").strip().lower().startswith("heading")


def _is_list(style_name: str, text: str) -> bool:
    lowered = (style_name or "").strip().lower()
    return lowered.startswith("list") or bool(re.match(r"^[\u2022\-*]\s+", text or ""))


def _clean(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").strip())


def _semantic_tags(title: str, text: str = "") -> list[str]:
    low = f"{title} {text}".lower()
    tags = []
    for needle in (
        "timeline",
        "governance",
        "solution",
        "scope",
        "methodology",
        "testing",
        "training",
        "assumption",
        "architecture",
        "migration",
        "security",
    ):
        if needle in low:
            tags.append(needle)
    return list(dict.fromkeys(tags))


def _iter_block_items(parent: DocxDocument | _Cell) -> Iterator[Paragraph | Table]:
    if isinstance(parent, DocxDocument):
        parent_elm = parent.element.body
    else:
        parent_elm = parent._tc
    for child in parent_elm.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def _extract_image_parts(doc: Document, template_id: str) -> dict[str, TemplateImage]:
    settings = get_settings()
    image_dir = settings.assets_path / "template-images" / template_id
    image_dir.mkdir(parents=True, exist_ok=True)

    images: dict[str, TemplateImage] = {}
    image_index = 0
    for rel_id, related in doc.part.related_parts.items():
        content_type = getattr(related, "content_type", "")
        if not content_type.startswith("image/"):
            continue
        ext = Path(getattr(related.partname, "filename", "")).suffix or ".png"
        filename = f"image_{image_index:03d}{ext}"
        asset_path = image_dir / filename
        asset_path.write_bytes(related.blob)
        images[rel_id] = TemplateImage(
            index=image_index,
            filename=filename,
            asset_path=str(asset_path),
            asset_url=f"/assets/template-images/{template_id}/{filename}",
            width=0.0,
            height=0.0,
            caption="",
            section="",
            page=0,
            purpose="template",
            semantic_tags=[],
        )
        image_index += 1
    return images


def _paragraph_image_rel_ids(paragraph: Paragraph) -> list[str]:
    rel_ids: list[str] = []
    for drawing in paragraph._p.xpath(".//a:blip"):
        embed = drawing.get(qn("r:embed"))
        if embed:
            rel_ids.append(embed)
    return rel_ids


def _static_section(title: str) -> bool:
    normalized = re.sub(r"[^a-z0-9]+", " ", (title or "").lower()).strip()
    return normalized in _STATIC_SECTIONS


def parse_template_docx(path: str | Path, proposal_family: str = "") -> TemplateDocumentArtifact:
    path = Path(path)
    template_id = uuid4().hex
    doc = Document(str(path))
    sections: list[TemplateSectionNode] = []
    blocks: list[TemplateBlock] = []
    stack: list[TemplateSectionNode] = []
    table_index = 0
    order = 0
    metadata: dict[str, Any] = {"paragraphs": len(doc.paragraphs), "tables": len(doc.tables)}
    image_parts = _extract_image_parts(doc, template_id)
    all_images: list[TemplateImage] = []
    current_section_title = "Document Overview"

    def current_target() -> TemplateSectionNode | None:
        return stack[-1] if stack else None

    for item in _iter_block_items(doc):
        if isinstance(item, Paragraph):
            text = _clean(item.text)
            style_name = getattr(getattr(item, "style", None), "name", "") or ""
            if _is_heading(style_name) and text:
                level_match = re.search(r"(\d+)$", style_name)
                level = int(level_match.group(1)) if level_match else 1
                node = TemplateSectionNode(title=text, level=level)
                while stack and stack[-1].level >= level:
                    stack.pop()
                if stack:
                    stack[-1].subsections.append(node)
                else:
                    sections.append(node)
                stack.append(node)
                current_section_title = text
                blocks.append(
                    TemplateBlock(
                        kind="heading",
                        section_title=text,
                        heading_level=level,
                        text=text,
                        style=style_name,
                        order=order,
                        static=_static_section(text),
                        editable=False,
                        adaptation_hint="Preserve the original heading unless explicitly renamed.",
                    )
                )
                order += 1
            elif text:
                target = current_target()
                if target is None:
                    target = TemplateSectionNode(title=current_section_title, level=1)
                    sections.append(target)
                    stack.append(target)
                target.paragraphs.append(TemplateParagraph(text=text, style=style_name, level=0))
                block_kind = "list" if _is_list(style_name, text) else "paragraph"
                list_items = [_clean(re.sub(r"^[\u2022\-*]\s+", "", text))] if block_kind == "list" else []
                blocks.append(
                    TemplateBlock(
                        kind=block_kind,
                        section_title=target.title,
                        heading_level=0,
                        text=text,
                        style=style_name,
                        order=order,
                        items=list_items,
                        static=_static_section(target.title),
                        editable=True,
                        adaptation_hint=(
                            "Adapt conservatively from the reference wording and preserve structure."
                            if not _static_section(target.title)
                            else "Keep this block unchanged unless explicitly edited by the user."
                        ),
                    )
                )
                order += 1

            target = current_target()
            for rel_id in _paragraph_image_rel_ids(item):
                image = image_parts.get(rel_id)
                if image is None:
                    continue
                image_copy = image.model_copy(deep=True)
                image_copy.section = target.title if target else current_section_title
                image_copy.semantic_tags = _semantic_tags(image_copy.section, text)
                if target:
                    target.images.append(image_copy)
                all_images.append(image_copy)
                blocks.append(
                    TemplateBlock(
                        kind="image",
                        section_title=image_copy.section,
                        order=order,
                        image=image_copy,
                        static=True,
                        editable=False,
                        adaptation_hint="Preserve or manually replace this reference image.",
                    )
                )
                order += 1
        else:
            target = current_target()
            if target is None:
                target = TemplateSectionNode(title=current_section_title, level=1)
                sections.append(target)
                stack.append(target)
            rows = [[_clean(cell.text) for cell in row.cells] for row in item.rows]
            rows = [row for row in rows if any(cell for cell in row)]
            table_node = TemplateTable(
                index=table_index,
                rows=len(rows),
                cols=max((len(row) for row in rows), default=0),
                style="",
                caption="",
            )
            target.tables.append(table_node)
            blocks.append(
                TemplateBlock(
                    kind="table",
                    section_title=target.title,
                    order=order,
                    table_rows=rows,
                    static=_static_section(target.title),
                    editable=True,
                    adaptation_hint="Preserve the table structure and edit only values that must change.",
                )
            )
            table_index += 1
            order += 1

    root = TemplateDocumentArtifact(
        template_id=template_id,
        name=path.stem,
        source_file=str(path),
        proposal_family=proposal_family or "General",
        sections=sections,
        blocks=blocks,
        images=all_images,
        metadata={
            **metadata,
            "block_count": len(blocks),
            "image_count": len(all_images),
            "top_level_sections": len(sections),
        },
    )
    return root


def default_template_artifact() -> TemplateDocumentArtifact:
    settings = get_settings()
    path = settings.proposal_template_path
    if path.exists():
        return parse_template_docx(path, proposal_family="Temenos")
    return TemplateDocumentArtifact(name="empty-template", proposal_family="General")
